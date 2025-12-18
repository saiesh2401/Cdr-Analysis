from docx import Document
import os

MASTER_SOURCE = "IP Letters latest.docx"

def extract_first_letter():
    doc = Document(MASTER_SOURCE)
    header_text = "OFFICE OF THE CYBER CRIME UNIT"
    
    # Locate split point (2nd header)
    count = 0
    split_element = None
    for p in doc.paragraphs:
        if header_text in p.text:
            count += 1
            if count == 2:
                split_element = p._element
                break
    
    # Truncate
    if split_element:
        body = doc.element.body
        children = list(body)
        deleting = False
        for child in children:
            if child == split_element:
                deleting = True
            if deleting:
                body.remove(child)
                
    doc.save("Master_Clean.docx")
    return "Master_Clean.docx"

def create_isp_template(master_path, filename, isp_to_block, subject_placeholder):
    doc = Document(master_path)
    
    # 1. Update "To" Block
    # Look for "The Nodal Officer" and the lines below it
    matched = False
    for i, p in enumerate(doc.paragraphs):
        if "The Nodal Officer" in p.text:
            # The next paragraph usually contains the ISP Name
            # We want to replace "Airtel" or whatever is there with our target ISP
            # Or better, just hardcode the block.
            # Let's find "Airtel" specifically in the doc and replace it if it's in the header area.
            pass

    # Brute force replace "Airtel" -> Target ISP Name (for the To Section)
    # Be careful not to replace it in the subject line yet (we'll overwrite subject line completely)
    
    # 2. Overwrite Subject Line
    for p in doc.paragraphs:
        if p.text.strip().startswith("Subject"):
            p.text = subject_placeholder
            break
            
    # 3. Replace "Airtel" in the body/To address with specific name
    # The source doc has "Airtel".
    for p in doc.paragraphs:
        if "Airtel" in p.text and "Subject" not in p.text: 
             p.text = p.text.replace("Airtel", isp_to_block)
             
    doc.save(filename)
    print(f"Created {filename}")

# Execution
clean_master = extract_first_letter()

subject_line = "Subject:- Reg provide information in case the case FIR No. {FIR_NO}, PS Special Cell ({EMAIL})-{ISP_NAME}"

# JIO
create_isp_template(clean_master, "JIO Template.docx", "Reliance Jio Infocomm Ltd.", subject_line)

# AIRTEL
create_isp_template(clean_master, "Airtel Template.docx", "Bharti Airtel Ltd.", subject_line)

# VI
create_isp_template(clean_master, "VI Template.docx", "Vodafone Idea Ltd.", subject_line)
