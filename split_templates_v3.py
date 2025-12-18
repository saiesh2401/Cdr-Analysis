from docx import Document
import copy

SOURCE = "IP Letters latest.docx"
SUBJECT_PLACEHOLDER = "Subject:- Reg provide information in case the case FIR No. {FIR_NO}, PS Special Cell ({EMAIL})-{ISP_NAME}"

def get_params():
    doc = Document(SOURCE)
    indices = [i for i, p in enumerate(doc.paragraphs) if "OFFICE OF THE CYBER CRIME UNIT" in p.text]
    # We expect indices for section 1 (0), 2 (ignore), 3 (VI), 4 (ignore), 5 (JIO)
    # The list is 0-indexed relative to paragraphs.
    # Actually the previously found indices were [1, 34, 70, 104, 138]
    return indices

def save_section(start_idx, end_idx, filename, isp_name):
    # Load fresh doc
    doc = Document(SOURCE)
    
    # We want to keep [start_idx, end_idx)
    # Strategy: 
    # 1. Remove everything AFTER end_idx (if end_idx is not None)
    # 2. Remove everything BEFORE start_idx
    
    # It's safer to remove from bottom to top to avoid index shifts, 
    # but removing range [0, start_idx] shifts everything up.
    
    # Let's iterate backwards specifically.
    
    total_len = len(doc.paragraphs)
    
    # 1. Delete after end
    if end_idx is not None:
        # Delete from end_idx to total_len
        # We need to find the elements.
        # This is tricky with python-docx because paragraphs don't cover tables.
        # Tables are siblings to paragraphs in doc.element.body
        pass
        
    # ALTERNATIVE: Use element parsing.
    doc = Document(SOURCE)
    body = doc.element.body
    
    # Find the paragraph elements corresponding to strict indices?
    # No, paragraphs list only contains paragraphs, skipping tables.
    # The header "OFFICE..." is a paragraph.
    
    # Let's count headers again using full body iteration to include tables logic if needed?
    # No, keep it simple. The headers define the breaks.
    
    # Precise method:
    # 1. Find the P element for Start Header.
    # 2. Find the P element for End Header (or End of Doc).
    # 3. Iterate body elements. keep flag False until Start Header. Set True. 
    #    If End Header hit, Set False.
    #    Remove elements where flag is False.
    
    header_text = "OFFICE OF THE CYBER CRIME UNIT"
    
    target_start_match = start_idx # 1st match, 3rd match, or 5th match
    target_end_match = end_idx     # match index to stop at
    
    current_match = 0
    keep_element = False
    
    elements_to_remove = []
    
    for child in body:
        # Check if this child is a paragraph with our header
        is_header = False
        if child.tag.endswith('p'):
            # It's a paragraph element, we can verify text potentially?
            # XML text extraction is annoying.
            # Let's assume the order is preserved.
            pass

    # RE-STRATEGY: 
    # The previous simple truncation worked because we only wanted the first one.
    # Now we want the middle ones.
    
    # Let's try a different approach:
    # Identify the P objects that are headers.
    headers = [p for p in doc.paragraphs if header_text in p.text]
    
    # headers[0] -> Airtel Start
    # headers[2] -> VI Start
    # headers[4] -> JIO Start
    
    start_node = headers[start_idx]._element
    end_node = headers[end_idx]._element if end_idx < len(headers) else None
    
    children = list(body)
    kept_count = 0
    
    in_zone = False
    
    for child in children:
        if child == start_node:
            in_zone = True
        
        if end_node and child == end_node:
            in_zone = False
            
        if not in_zone:
            body.remove(child)
        else:
            kept_count += 1
            
    # Now Update Subject
    for p in doc.paragraphs:
        if p.text.strip().startswith("Subject"):
            p.text = SUBJECT_PLACEHOLDER
            # Also replace any stray "Airtel" in body if this is JIO/VI?
            # User said templates are separate, so they should be correct already.
            # But the Subject needs to be dynamic.
            break
            
    doc.save(filename)
    print(f"Saved {filename} with {kept_count} elements.")

# Indices of headers: 0, 1, 2, 3, 4
# Airtel: Start 0, End 1 (Start of next garbage letter)
# VI: Start 2, End 3
# JIO: Start 4, End None

try:
    print("--- Splitting ---")
    # Airtel
    save_section(0, 1, "Airtel Template.docx", "Bharti Airtel Ltd.")
    
    # VI
    save_section(2, 3, "VI Template.docx", "Vodafone Idea Ltd.")
    
    # JIO
    save_section(4, 99, "JIO Template.docx", "Reliance Jio Infocomm Ltd.") 
    
except Exception as e:
    print(f"Error: {e}")
