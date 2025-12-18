import os
import json
import datetime
from bs4 import BeautifulSoup
from ipwhois import IPWhois
import openpyxl
from docx import Document
from collections import defaultdict
import warnings

import sys

# Suppress warnings
warnings.filterwarnings("ignore")

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

class ISPProcessor:
    def __init__(self, output_dir="Generated_Letters", cache_file="isp_cache.json"):
        self.output_dir = output_dir
        # Cache file should live in the actual user directory, not the temp bundle, 
        # so we DO NOT use resource_path for it (unless we want read-only cache).
        self.cache_file = cache_file 
        self.isp_cache = self._load_cache()
        
        # Ensure output dir exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _load_cache(self):
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def _save_cache(self):
        try:
            with open(self.cache_file, 'w') as f:
                json.dump(self.isp_cache, f)
        except Exception as e:
            print(f"Failed to save cache: {e}")

    def get_isp(self, ip, log_callback=None):
        if ip in self.isp_cache:
            return self.isp_cache[ip]
        
        try:
            obj = IPWhois(ip)
            res = obj.lookup_rdap()
            org = res.get('network', {}).get('name', 'Unknown')
            if not org:
                org = res.get('asn_description', 'Unknown')
            
            if log_callback:
                log_callback(f"Lookup: {ip} -> {org}")
            
            org_upper = org.upper()
            if "JIO" in org_upper or "RELIANCE" in org_upper:
                isp = "JIO"
            elif "AIRTEL" in org_upper or "BHARTI" in org_upper:
                isp = "AIRTEL"
            elif "VODAFONE" in org_upper or "VI" in org_upper or "IDEA" in org_upper:
                isp = "VI"
            elif "BSNL" in org_upper:
                isp = "BSNL"
            else:
                isp = "OTHER"
                
            self.isp_cache[ip] = isp
            self._save_cache()
            return isp
        except Exception as e:
            if log_callback:
                log_callback(f"Error looking up {ip}: {e}")
            return "Unknown"

    def parse_html(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"HTML file not found: {file_path}")

        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')

        data = []
        metadata = {"name": "Unknown", "email": "Unknown"}
        
        # Extract Metadata (Name, Email)
        # Usually in the first few sections
        list_items = soup.find_all('li')
        for li in list_items:
            text = li.get_text().strip()
            if text.startswith("Name:"):
                metadata["name"] = text.replace("Name:", "").strip()
            elif text.startswith("e-Mail:"):
                metadata["email"] = text.replace("e-Mail:", "").strip()

        tables = soup.find_all('table')
        for table in tables:
            headers = [th.get_text().strip() for th in table.find_all('th')]
            if "IP Address" in headers and "Timestamp" in headers:
                rows = table.find_all('tr')[1:] 
                for row in rows:
                    cols = row.find_all('td')
                    if len(cols) >= 2:
                        ts = cols[0].get_text().strip()
                        ip = cols[1].get_text().strip()
                        if ip:
                             data.append({'timestamp': ts, 'ip': ip})
        return metadata, data

    def process_data(self, raw_data, progress_callback=None, log_callback=None):
        grouped = defaultdict(list)
        total = len(raw_data)
        
        for idx, entry in enumerate(raw_data):
            if progress_callback:
                progress_callback(idx + 1, total)
                
            ip = entry['ip']
            ts_str = entry['timestamp']
            
            try:
                # Format is "YYYY-MM-DD HH:MM:SS Z", which is UTC.
                # Remove Z, parse time, and add 5h 30m for IST.
                dt_utc = datetime.datetime.strptime(ts_str.replace(" Z", ""), "%Y-%m-%d %H:%M:%S")
                dt = dt_utc + datetime.timedelta(hours=5, minutes=30)
            except:
                dt = datetime.datetime.now()
                
            isp = self.get_isp(ip, log_callback)
            entry['datetime'] = dt
            grouped[isp].append(entry)
            
        return grouped

    def fill_jio_excel(self, entries, template_path, output_name="JIO_Data.xlsx"):
        template_path = resource_path(template_path)
        if not os.path.exists(template_path):
            return False, "Template missing"

        wb = openpyxl.load_workbook(template_path)
        sheet = wb.active
        current_row = sheet.max_row + 1
        
        for item in entries:
            ip = item['ip']
            dt = item['datetime']
            from_dt = dt - datetime.timedelta(minutes=5)
            to_dt = dt + datetime.timedelta(minutes=5)
            ip_type = "IPV6" if ":" in ip else "IPV4"
            
            sheet.cell(row=current_row, column=1, value=ip_type)
            sheet.cell(row=current_row, column=2, value=ip)
            sheet.cell(row=current_row, column=3, value=from_dt.strftime("%Y%m%d"))
            sheet.cell(row=current_row, column=4, value=from_dt.strftime("%H%M%S"))
            sheet.cell(row=current_row, column=5, value=to_dt.strftime("%Y%m%d"))
            sheet.cell(row=current_row, column=6, value=to_dt.strftime("%H%M%S"))
            current_row += 1
            
        out_path = os.path.join(self.output_dir, output_name)
        wb.save(out_path)
        return True, out_path

    def fill_jio_txt(self, entries, output_name="JIO_Data.txt"):
        out_path = os.path.join(self.output_dir, output_name)
        
        # Header from sample file
        # Note: Sample had "From Time" twice. I will keep it exactly as sample unless directed otherwise,
        # but standard logic usually suggests To Time. I'll stick to the sample header text to be safe.
        header = "Type\tSearch Value\tFrom Date YYYYMMDD\tFrom Time HHMMSS (IST)\tTo Date YYYYMMDD\tFrom Time HHMMSS (IST)\n"
        
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(header)
                
                for item in entries:
                    ip = item['ip']
                    dt = item['datetime']
                    from_dt = dt - datetime.timedelta(minutes=5)
                    to_dt = dt + datetime.timedelta(minutes=5)
                    ip_type = "IPV6" if ":" in ip else "IPV4"
                    
                    # JIO Format: YYYYMMDD, HHMMSS
                    row = [
                        ip_type,
                        ip,
                        from_dt.strftime("%Y%m%d"),
                        from_dt.strftime("%H%M%S"),
                        to_dt.strftime("%Y%m%d"),
                        to_dt.strftime("%H%M%S")
                    ]
                    f.write("\t".join(row) + "\n")
                    
            return True, out_path
        except Exception as e:
            return False, str(e)

    def fill_airtel_excel(self, entries, template_path, output_name="Airtel_Data.xlsx"):
        template_path = resource_path(template_path)
        if not os.path.exists(template_path):
             return False, "Template missing"

        wb = openpyxl.load_workbook(template_path)
        sheet = wb.active
        current_row = sheet.max_row + 1
        
        for item in entries:
            ip = item['ip']
            dt = item['datetime']
            ip_type = "IPV6" if ":" in ip else "IPV4"
            
            sheet.cell(row=current_row, column=1, value=ip_type)
            sheet.cell(row=current_row, column=2, value=ip)
            sheet.cell(row=current_row, column=3, value=dt.strftime("%d-%b-%Y"))
            sheet.cell(row=current_row, column=4, value=dt.strftime("%H:%M:%S"))
            current_row += 1

        out_path = os.path.join(self.output_dir, output_name)
        wb.save(out_path)
        return True, out_path

    def fill_generic_excel(self, entries, output_name="ISP_Data.xlsx"):
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = "IP Data"
        
        # Headers
        headers = ["IP Type", "IP Address", "Date", "Time", "From Date", "To Date"]
        for col_num, header in enumerate(headers, 1):
            sheet.cell(row=1, column=col_num, value=header)
            
        current_row = 2
        for item in entries:
            ip = item['ip']
            dt = item['datetime']
            from_dt = dt - datetime.timedelta(minutes=5)
            to_dt = dt + datetime.timedelta(minutes=5)
            ip_type = "IPV6" if ":" in ip else "IPV4"
            
            sheet.cell(row=current_row, column=1, value=ip_type)
            sheet.cell(row=current_row, column=2, value=ip)
            sheet.cell(row=current_row, column=3, value=dt.strftime("%d-%m-%Y"))
            sheet.cell(row=current_row, column=4, value=dt.strftime("%H:%M:%S"))
            sheet.cell(row=current_row, column=5, value=from_dt.strftime("%d-%m-%Y %H:%M:%S"))
            sheet.cell(row=current_row, column=6, value=to_dt.strftime("%d-%m-%Y %H:%M:%S"))
            current_row += 1
            
        out_path = os.path.join(self.output_dir, output_name)
        wb.save(out_path)
        return True, out_path

    def fill_word_letter(self, entries, isp_name, template_path, output_name=None, metadata=None):
        template_path = resource_path(template_path)
        if not os.path.exists(template_path):
            return False, "Template missing"

        if output_name is None:
            output_name = f"{isp_name}_Request_Letter.docx"

        doc = Document(template_path)
        
        # Map simple ISP code to Full Name for the placeholder
        isp_full_names = {
            "JIO": "Reliance Jio Infocomm Ltd.",
            "AIRTEL": "Bharti Airtel Ltd.",
            "VI": "Vodafone Idea Ltd."
        }
        full_isp_name = isp_full_names.get(isp_name, isp_name)

        # Replace Placeholders in Paragraphs
        if metadata:
            name = metadata.get("name", "")
            email = metadata.get("email", "")
            fir_no = metadata.get("fir_no", "")
            fir_date = metadata.get("fir_date", "")
            current_date = datetime.datetime.now().strftime("%d.%m.%Y")
            
            replacements = {
                "{NAME}": name,
                "{EMAIL}": email,
                "{FIR_NO}": fir_no,
                "{FIR_DATE}": fir_date,
                "{ISP_NAME}": full_isp_name,
                "{DATE}": current_date
            }

            for p in doc.paragraphs:
                for key, val in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, str(val))
        
        target_table = None
        for table in doc.tables:
            if not table.rows: continue
            headers = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
            if any("Search Value" in h for h in headers) or any("IP" in h for h in headers):
                target_table = table
                break
                
        if target_table:
            # CLEAR SAMPLE DATA: Remove all rows except the header
            # Iterate backwards to avoid index issues
            for i in range(len(target_table.rows) - 1, 0, -1):
                row = target_table.rows[i]
                # Remove XML element
                row._element.getparent().remove(row._element)

            # Now populate with new data
            for item in entries:
                ip = item['ip']
                dt = item['datetime']
                from_dt = dt - datetime.timedelta(minutes=5)
                to_dt = dt + datetime.timedelta(minutes=5)
                ip_type = "IPV6" if ":" in ip else "IPV4"

                # Create new row
                row = target_table.add_row()
                
                if len(row.cells) < len(target_table.columns):
                     # Should rarely happen with add_row() on uniform table
                     pass
                
                cells = row.cells
                
                # Define Format based on ISP
                if isp_name == "JIO":
                    d_fmt = "%Y%m%d"
                    t_fmt = "%H%M%S"
                elif isp_name == "VI":
                    # Image showed dots: 22.09.2025
                    d_fmt = "%d.%m.%Y" 
                    t_fmt = "%H:%M:%S"
                elif isp_name == "AIRTEL":
                    # User requested: 24/Jan/2025
                    d_fmt = "%d/%b/%Y"
                    t_fmt = "%H:%M:%S"
                else:
                    d_fmt = "%d-%m-%Y"
                    t_fmt = "%H:%M:%S"

                fmt_combined = f"{d_fmt}\n{t_fmt}"
                
                if len(cells) >= 6:
                    cells[0].text = ip_type
                    cells[1].text = ip
                    cells[2].text = from_dt.strftime(d_fmt) 
                    cells[3].text = from_dt.strftime(t_fmt)
                    cells[4].text = to_dt.strftime(d_fmt)
                    cells[5].text = to_dt.strftime(t_fmt)
                elif len(cells) >= 4:
                     cells[0].text = ip_type
                     cells[1].text = ip
                     cells[2].text = from_dt.strftime(fmt_combined)
                     cells[3].text = to_dt.strftime(fmt_combined)
                
            out_path = os.path.join(self.output_dir, output_name)
            doc.save(out_path)
            return True, out_path
        else:
            return False, "Table not found"
